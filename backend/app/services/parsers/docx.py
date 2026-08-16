"""python-docx DOCX 解析器（T051 / quickstart 解析边界）。

- 先以 zipfile 审查归档：条目数与解压后总大小上限（压缩炸弹/洪泛防护）；
- python-docx 只解析 OOXML 文本；不执行宏（vbaProject 内容不被解释）、
  外部链接或嵌入对象；含 ``../`` 成员名不会写出任何文件；
- 解析失败统一映射 ``20001`` 固定安全提示。
"""

import io
import zipfile
from importlib.metadata import version

from docx import Document as DocxDocument

from app.services.parsers.base import ParseError, ParseOutput, parse_failed

MAX_ZIP_ENTRIES = 1024


class DocxParser:
    name = "python-docx"

    def __init__(self) -> None:
        try:
            self.parser_version = version("python-docx")
        except Exception:  # noqa: BLE001
            self.parser_version = "unknown"

    def parse(self, content: bytes, max_expanded_bytes: int | None = None) -> ParseOutput:
        buffer = io.BytesIO(content)
        try:
            with zipfile.ZipFile(buffer) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_ZIP_ENTRIES:
                    raise parse_failed()
                expanded = sum(entry.file_size for entry in entries)
                if max_expanded_bytes is not None and expanded > max_expanded_bytes:
                    raise parse_failed()
        except (zipfile.BadZipFile, EOFError):
            raise parse_failed() from None
        except ParseError:
            raise
        except Exception:
            raise parse_failed() from None

        try:
            buffer.seek(0)
            document = DocxDocument(buffer)
        except Exception:
            raise parse_failed() from None

        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.extend(cell.text for cell in row.cells)
        return ParseOutput(
            normalized_text="\n".join(paragraphs),
            parser_name=self.name,
            parser_version=self.parser_version,
        )
